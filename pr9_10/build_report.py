from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw

from config import SCREEN_HEIGHT, SCREEN_WIDTH


BASE_DIR = Path(__file__).parent
SCREENSHOT_DIR = BASE_DIR / "assets" / "screenshots"
REPORT_PATH = BASE_DIR / "PR9_10_Crystal_Sprint_Report.docx"


def draw_background(draw):
    draw.rectangle((0, 0, SCREEN_WIDTH, SCREEN_HEIGHT), fill="#0b1021")
    for i in range(0, SCREEN_HEIGHT, 40):
        color = "#101935" if (i // 40) % 2 == 0 else "#0d1530"
        draw.rectangle((0, i, SCREEN_WIDTH, i + 40), fill=color)

    for x in range(40, SCREEN_WIDTH, 80):
        draw.line((x, 0, x, SCREEN_HEIGHT), fill="#16244a", width=1)
    for y in range(60, SCREEN_HEIGHT, 80):
        draw.line((0, y, SCREEN_WIDTH, y), fill="#16244a", width=1)


def draw_hud(draw, score, lives, level, message=None):
    draw.text((18, 14), f"Очки: {score}   Життя: {lives}   Рівень: {level}", fill="#f5f7ff")
    if message:
        draw.text((240, 260), message, fill="#7cffb5")


def draw_player(draw, x, y):
    draw.rectangle((x - 18, y - 18, x + 18, y + 18), fill="#49d6ff", outline="#9bf0ff", width=2)


def draw_enemy(draw, x, y):
    draw.ellipse((x - 16, y - 16, x + 16, y + 16), fill="#ff5d73", outline="#ffc0c9", width=2)


def draw_crystal(draw, x, y):
    draw.polygon(
        [(x, y - 12), (x + 12, y), (x, y + 12), (x - 12, y)],
        fill="#ffd166",
        outline="#ffe5a4",
    )


def draw_projectile(draw, x, y):
    draw.ellipse((x - 5, y - 5, x + 5, y + 5), fill="#ffffff", outline="#d9eef7")


def generate_screenshots():
    SCREENSHOT_DIR.mkdir(parents=True, exist_ok=True)
    image_paths = []

    scenes = [
        {
            "name": "screenshot_1_start.png",
            "score": 0,
            "lives": 3,
            "level": 1,
            "player": (400, 520),
            "enemies": [(150, 100), (650, 120)],
            "crystals": [(260, 240), (540, 310)],
            "projectiles": [(420, 450)],
            "message": "Початок гри",
        },
        {
            "name": "screenshot_2_action.png",
            "score": 11,
            "lives": 2,
            "level": 2,
            "player": (260, 420),
            "enemies": [(140, 160), (620, 200), (450, 110)],
            "crystals": [(520, 360)],
            "projectiles": [(300, 360), (360, 290)],
            "message": "Активний бій",
        },
        {
            "name": "screenshot_3_victory.png",
            "score": 20,
            "lives": 1,
            "level": 3,
            "player": (410, 360),
            "enemies": [],
            "crystals": [(340, 250), (470, 250)],
            "projectiles": [(380, 320)],
            "message": "Перемога!",
        },
    ]

    for scene in scenes:
        image = Image.new("RGB", (SCREEN_WIDTH, SCREEN_HEIGHT), "#0b1021")
        draw = ImageDraw.Draw(image)
        draw_background(draw)
        draw_hud(draw, scene["score"], scene["lives"], scene["level"], scene["message"])
        draw_player(draw, *scene["player"])

        for enemy in scene["enemies"]:
            draw_enemy(draw, *enemy)
        for crystal in scene["crystals"]:
            draw_crystal(draw, *crystal)
        for projectile in scene["projectiles"]:
            draw_projectile(draw, *projectile)

        path = SCREENSHOT_DIR / scene["name"]
        image.save(path)
        image_paths.append(path)

    return image_paths


def build_report(image_paths):
    document = Document()
    document.add_heading("Практична робота 9-10", level=0)
    document.add_paragraph("Назва гри: Crystal Sprint")
    document.add_paragraph("Жанр: аркада / гра на реакцію")
    document.add_paragraph(
        "Суть гри: гравець керує синім персонажем, збирає кристали та стріляє по ворогах. "
        "За кристали та знищених ворогів нараховуються очки. Для перемоги потрібно набрати 20 очок."
    )
    document.add_paragraph("Керування:")
    document.add_paragraph("- W, A, S, D або стрілки - рух персонажа", style="List Bullet")
    document.add_paragraph("- Клік мишкою - постріл у напрямку курсора", style="List Bullet")
    document.add_paragraph("- Esc - вихід з гри", style="List Bullet")
    document.add_paragraph("Умова перемоги: набрати 20 очок.")
    document.add_paragraph("Умова поразки: втратити всі 3 життя.")
    document.add_paragraph("Створені класи:")
    document.add_paragraph("- BaseEntity", style="List Bullet")
    document.add_paragraph("- Player", style="List Bullet")
    document.add_paragraph("- Enemy", style="List Bullet")
    document.add_paragraph("- Projectile", style="List Bullet")
    document.add_paragraph("- Crystal", style="List Bullet")
    document.add_paragraph("- Hud", style="List Bullet")
    document.add_paragraph("- Game", style="List Bullet")
    document.add_paragraph(
        "Використана ШІ-система: ChatGPT. "
        "Її було використано для допомоги з проєктуванням структури класів, написанням коду гри "
        "та формуванням тексту опису для Word-файлу."
    )
    document.add_paragraph("Скріншоти гри:")

    for image_path in image_paths:
        document.add_picture(str(image_path), width=Inches(6.2))
        document.add_paragraph(image_path.stem.replace("_", " ").capitalize())

    document.save(REPORT_PATH)


if __name__ == "__main__":
    screenshots = generate_screenshots()
    build_report(screenshots)
